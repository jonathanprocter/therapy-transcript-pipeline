import os
import logging
import io
from typing import Dict, Optional, Tuple
import PyPDF2
import docx
from datetime import datetime
import re

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Service for processing different document types"""
    
    def __init__(self):
        self.supported_types = {
            '.pdf': self._process_pdf,
            '.txt': self._process_txt,
            '.docx': self._process_docx
        }
    
    def process_document(self, file_content: bytes, filename: str) -> Dict:
        """Process a document and extract text content"""
        try:
            file_ext = os.path.splitext(filename)[1].lower()
            
            if file_ext not in self.supported_types:
                raise ValueError(f"Unsupported file type: {file_ext}")
            
            # Process the document based on its type
            processor = self.supported_types[file_ext]
            text_content = processor(file_content)
            
            if not text_content or len(text_content.strip()) == 0:
                raise ValueError("No text content extracted from document")
            
            # Extract metadata and clean content
            processed_data = {
                'raw_content': text_content,
                'cleaned_content': self._clean_text(text_content),
                'word_count': len(text_content.split()),
                'character_count': len(text_content),
                'file_type': file_ext[1:],  # Remove the dot
                'extracted_date': self._extract_session_date(text_content),
                'processing_metadata': {
                    'processed_at': datetime.utcnow().isoformat(),
                    'processor_version': '1.0',
                    'extraction_method': file_ext[1:]
                }
            }
            
            logger.info(f"Successfully processed {filename}: {processed_data['word_count']} words extracted")
            return processed_data
            
        except Exception as e:
            logger.error(f"Error processing document {filename}: {str(e)}")
            raise
    
    def _process_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = ""
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += f"\n--- Page {page_num + 1} ---\n{page_text}"
                except Exception as e:
                    logger.warning(f"Error extracting text from PDF page {page_num + 1}: {str(e)}")
                    continue
            
            if not text_content.strip():
                raise ValueError("No text could be extracted from PDF")
            
            return text_content
            
        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}")
            raise
    
    def _process_txt(self, file_content: bytes) -> str:
        """Extract text from TXT file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text_content = file_content.decode(encoding)
                    logger.info(f"Successfully decoded TXT file using {encoding} encoding")
                    return text_content
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Could not decode text file with any supported encoding")
            
        except Exception as e:
            logger.error(f"Error processing TXT file: {str(e)}")
            raise
    
    def _process_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            docx_file = io.BytesIO(file_content)
            doc = docx.Document(docx_file)
            
            text_content = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content += cell.text + " "
                    text_content += "\n"
            
            if not text_content.strip():
                raise ValueError("No text could be extracted from DOCX file")
            
            return text_content
            
        except Exception as e:
            logger.error(f"Error processing DOCX file: {str(e)}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove page markers
        text = re.sub(r'--- Page \d+ ---', '', text)
        
        # Remove non-printable characters except newlines and tabs
        text = ''.join(char for char in text if char.isprintable() or char in '\n\t')
        
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove excessive newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()
    
    def _extract_session_date(self, text: str) -> Optional[datetime]:
        """Extract session date from text content"""
        try:
            # Common date patterns in therapy transcripts
            date_patterns = [
                r'(?:Session|Date|Meeting):\s*(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})',
                r'(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})',
                r'(\w+\s+\d{1,2},?\s+\d{4})',  # "January 15, 2024"
                r'(\d{4}-\d{2}-\d{2})',  # ISO format
            ]
            
            for pattern in date_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    date_str = matches[0]
                    
                    # Try to parse the date
                    for date_format in ['%m/%d/%Y', '%m-%d-%Y', '%m/%d/%y', '%m-%d-%y', 
                                       '%B %d, %Y', '%B %d %Y', '%Y-%m-%d']:
                        try:
                            return datetime.strptime(date_str, date_format)
                        except ValueError:
                            continue
            
            logger.info("No recognizable date found in document")
            return None
            
        except Exception as e:
            logger.warning(f"Error extracting session date: {str(e)}")
            return None
    
    def extract_client_identifier(self, text: str, filename: str) -> Optional[str]:
        """Extract client identifier from text or filename"""
        try:
            # Try to extract from filename first
            filename_parts = os.path.splitext(filename)[0].split('_')
            if len(filename_parts) > 1:
                # Assume first part might be client identifier
                potential_id = filename_parts[0]
                if potential_id and len(potential_id) > 2:
                    return potential_id
            
            # Try to extract from text content
            client_patterns = [
                r'(?:Client|Patient):\s*([A-Za-z\s]+)',
                r'(?:Name|Client ID):\s*([A-Za-z0-9\s]+)',
            ]
            
            for pattern in client_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    client_id = matches[0].strip()
                    if len(client_id) > 2:
                        return client_id
            
            # Fallback: use filename without extension
            return os.path.splitext(filename)[0]
            
        except Exception as e:
            logger.warning(f"Error extracting client identifier: {str(e)}")
            return os.path.splitext(filename)[0]
    
    def validate_content(self, text: str) -> Dict:
        """Validate if content appears to be a therapy transcript"""
        validation_result = {
            'is_valid': True,
            'confidence': 0.0,
            'issues': [],
            'characteristics': []
        }
        
        if not text or len(text.strip()) < 100:
            validation_result['is_valid'] = False
            validation_result['issues'].append("Content too short to be a meaningful transcript")
            return validation_result
        
        # Check for therapy-related keywords
        therapy_keywords = [
            'session', 'therapy', 'therapist', 'client', 'patient', 'counseling',
            'feelings', 'thoughts', 'emotions', 'discussion', 'talk', 'share'
        ]
        
        text_lower = text.lower()
        keyword_count = sum(1 for keyword in therapy_keywords if keyword in text_lower)
        
        if keyword_count >= 3:
            validation_result['characteristics'].append("Contains therapy-related terminology")
            validation_result['confidence'] += 0.3
        
        # Check for dialogue patterns
        dialogue_patterns = [
            r'\b(therapist|counselor|doctor):\s',
            r'\b(client|patient):\s',
            r'\b(t:|c:|p:)\s',
            r'"[^"]*"',  # Quoted speech
        ]
        
        dialogue_count = 0
        for pattern in dialogue_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                dialogue_count += 1
        
        if dialogue_count >= 2:
            validation_result['characteristics'].append("Contains dialogue structure")
            validation_result['confidence'] += 0.4
        
        # Check word count (therapy sessions typically 1000+ words)
        word_count = len(text.split())
        if word_count >= 1000:
            validation_result['characteristics'].append("Adequate length for therapy session")
            validation_result['confidence'] += 0.2
        elif word_count < 500:
            validation_result['issues'].append("Content may be too short for a full therapy session")
            validation_result['confidence'] -= 0.1
        
        # Check for time markers
        time_patterns = [
            r'\d{1,2}:\d{2}',  # Time stamps
            r'(?:beginning|middle|end) of session',
            r'(?:minutes|hour) into',
        ]
        
        for pattern in time_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                validation_result['characteristics'].append("Contains time references")
                validation_result['confidence'] += 0.1
                break
        
        # Final validation
        validation_result['confidence'] = max(0.0, min(1.0, validation_result['confidence']))
        if validation_result['confidence'] < 0.3:
            validation_result['is_valid'] = False
            validation_result['issues'].append("Low confidence that this is a therapy transcript")
        
        return validation_result
